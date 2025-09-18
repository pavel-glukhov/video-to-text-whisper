import os
import shutil
import subprocess
import tempfile
import whisper
import torch
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

VIDEO_EXTENSIONS = (".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv")

FONT_PATH = os.path.join(os.path.dirname(__file__), "..", "fonts", "DejaVuSans.ttf")
PDF_FONT = "Helvetica"


def register_font():
    """
    Register DejaVuSans font for PDF if available, otherwise fallback to Helvetica.
    """
    global PDF_FONT
    if os.path.exists(FONT_PATH):
        try:
            pdfmetrics.registerFont(TTFont("DejaVuSans", FONT_PATH))
            PDF_FONT = "DejaVuSans"
            print(f"[INFO] Registered font: DejaVuSans ({FONT_PATH})")
        except Exception as e:
            print(f"[WARNING] Failed to register DejaVuSans, fallback to Helvetica: {e}")
    else:
        print(f"[WARNING] Font not found at {FONT_PATH}, fallback to Helvetica")


register_font()


def get_device(device_preference: str = None) -> str:
    """
    Detects which device should be used (CPU or CUDA).
    """
    if device_preference == "cuda" and torch.cuda.is_available():
        print("[INFO] Using CUDA (GPU)")
        return "cuda"
    elif device_preference == "cpu":
        print("[INFO] Using CPU (forced)")
        return "cpu"
    device = "cuda" if torch.cuda.is_available() else "cpu"
    print(f"[INFO] Auto-detected device: {device.upper()}")
    return device


def extract_audio(video_path: str, audio_path: str, ffmpeg_path: str) -> str:
    """
    Extract audio from video using ffmpeg (16kHz mono WAV).
    """
    print(f"[INFO] Extracting audio from: {video_path}")
    command = [
        ffmpeg_path, "-y", "-i", video_path,
        "-ar", "16000", "-ac", "1", audio_path
    ]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    print(f"[INFO] Audio saved to: {audio_path}")
    return audio_path


def transcribe_with_whisper(audio_path: str, model_name: str = "base", device: str = "cpu", block_seconds: int = 60) -> list[str]:
    """
    Transcribe audio using Whisper and group text into time blocks.
    """
    print(f"[INFO] Loading Whisper model: {model_name} on {device.upper()}")
    model = whisper.load_model(model_name, device=device)

    print("[INFO] Starting transcription...")
    result = model.transcribe(audio_path, verbose=False)

    blocks, current_text, block_start, last_end = [], [], None, None
    for seg in result["segments"]:
        start, end, text = seg["start"], seg["end"], seg["text"].strip()
        if block_start is None:
            block_start = start
        current_text.append(text)
        last_end = end
        if last_end - block_start >= block_seconds:
            blocks.append(f"[{block_start:.2f} - {last_end:.2f}] {' '.join(current_text)}")
            current_text, block_start = [], None
    if current_text:
        blocks.append(f"[{block_start:.2f} - {last_end:.2f}] {' '.join(current_text)}")

    print(f"[INFO] Transcription complete. Total blocks: {len(blocks)}")
    return blocks


def save_to_docx(blocks: list[str], output_path: str):
    """
    Save transcription blocks into a DOCX file.
    """
    print(f"[INFO] Saving DOCX to: {output_path}")
    doc = Document()
    doc.add_heading("Video Transcription", level=1)
    for block in blocks:
        doc.add_paragraph(block)
    doc.save(output_path)


def save_to_txt(blocks: list[str], output_path: str):
    """
    Save transcription blocks into a TXT file (UTF-8).
    """
    print(f"[INFO] Saving TXT to: {output_path}")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("Video Transcription\n\n")
        for block in blocks:
            f.write(block + "\n")


def save_to_pdf(blocks: list[str], output_path: str):
    """
    Save transcription blocks into a PDF file with proper font support.
    """
    print(f"[INFO] Saving PDF to: {output_path}")
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Cyrillic", fontName=PDF_FONT, fontSize=12, leading=14))
    styles.add(ParagraphStyle(name="CyrillicTitle", fontName=PDF_FONT, fontSize=16, leading=18, spaceAfter=12))

    story = [Paragraph("Video Transcription", styles["CyrillicTitle"]), Spacer(1, 12)]
    for block in blocks:
        story.append(Paragraph(block, styles["Cyrillic"]))
        story.append(Spacer(1, 8))

    doc.build(story)
    print(f"[INFO] PDF saved: {output_path}")


def process_video(video_path: str, model_name: str, device: str, block_seconds: int, formats: list[str], ffmpeg_path: str):
    """
    Process a single video: extract audio, transcribe, and save in requested formats.
    """
    print(f"\n=== Processing video: {video_path} ===")
    temp_dir = tempfile.mkdtemp(prefix="whisper_audio_")
    audio_file = os.path.join(temp_dir, "temp_audio.wav")
    try:
        extract_audio(video_path, audio_file, ffmpeg_path)
        blocks = transcribe_with_whisper(audio_file, model_name=model_name, device=device, block_seconds=block_seconds)

        base_path = os.path.splitext(video_path)[0]
        for fmt in formats:
            if fmt == "docx":
                save_to_docx(blocks, base_path + ".docx")
            elif fmt == "txt":
                save_to_txt(blocks, base_path + ".txt")
            elif fmt == "pdf":
                save_to_pdf(blocks, base_path + ".pdf")

        print(f"[INFO] Finished processing: {video_path}")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
