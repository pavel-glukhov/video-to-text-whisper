import os
import sys
import shutil
import subprocess
import tempfile
import whisper
from docx import Document
import torch
import warnings
import argparse

warnings.filterwarnings("ignore", category=UserWarning, module="whisper")

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
CACHE_DIR = os.path.join(PROJECT_DIR, "cache")
os.environ["XDG_CACHE_HOME"] = CACHE_DIR

VIDEO_EXTENSIONS = (".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv")


def get_device(device_preference: str = None) -> str:
    """
    Determines which device to use: CPU or GPU (CUDA).

    Args:
        device_preference (str, optional): "cuda" or "cpu". Defaults to None, auto-detect.

    Returns:
        str: Device to use.
    """
    if device_preference:
        if device_preference.lower() == "cuda" and torch.cuda.is_available():
            print(f"[INFO] Using GPU (forced): {torch.cuda.get_device_name(0)}")
            return "cuda"
        elif device_preference.lower() == "cpu":
            print("[INFO] Using CPU (forced)")
            return "cpu"
        else:
            print(f"[WARNING] Requested device '{device_preference}' not available. Auto-detecting...")
    if torch.cuda.is_available():
        print(f"[INFO] Using GPU: {torch.cuda.get_device_name(0)}")
        return "cuda"
    print("[INFO] CUDA not found, using CPU")
    return "cpu"


def get_ffmpeg_path() -> str:
    system_ffmpeg = shutil.which("ffmpeg")
    if system_ffmpeg:
        print(f"[INFO] Using system ffmpeg: {system_ffmpeg}")
        return system_ffmpeg
    local_ffmpeg = os.path.join(PROJECT_DIR, "ffmpeg", "bin", "ffmpeg.exe")
    if os.path.exists(local_ffmpeg):
        print(f"[WARNING] System ffmpeg not found, using local: {local_ffmpeg}")
        return local_ffmpeg
    print("[ERROR] ffmpeg not found! Install globally or place in ./ffmpeg/bin/")
    sys.exit(1)


def extract_audio(video_path: str, audio_path: str) -> str:
    command = [
        FFMPEG_PATH,
        "-y",
        "-i", video_path,
        "-ar", "16000",
        "-ac", "1",
        audio_path
    ]
    subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return audio_path


def transcribe_with_whisper(audio_path: str, model_name: str = "base", block_seconds: int = 60) -> list[str]:
    print(f"Loading model '{model_name}' on {DEVICE}...")
    model = whisper.load_model(model_name, device=DEVICE)

    print("Starting transcription...")
    result = model.transcribe(audio_path, verbose=False)

    blocks = []
    current_text = []
    block_start = None
    last_end = None

    for seg in result["segments"]:
        start = seg["start"]
        end = seg["end"]
        text = seg["text"].strip()

        if block_start is None:
            block_start = start

        current_text.append(text)
        last_end = end

        if last_end - block_start >= block_seconds:
            block_text = " ".join(current_text)
            blocks.append(f"[{block_start:.2f} - {last_end:.2f}] {block_text}")
            current_text = []
            block_start = None

    if current_text:
        block_text = " ".join(current_text)
        blocks.append(f"[{block_start:.2f} - {last_end:.2f}] {block_text}")

    return blocks


def save_to_docx(blocks: list[str], output_path: str):
    doc = Document()
    doc.add_heading("Video Transcription", level=1)
    for block in blocks:
        doc.add_paragraph(block)
    doc.save(output_path)


def process_video(video_path: str, model_name: str = "base", block_seconds: int = 60):
    print(f"\n=== Processing: {video_path} ===")
    temp_dir = tempfile.mkdtemp(prefix="whisper_audio_")
    audio_file = os.path.join(temp_dir, "temp_audio.wav")
    docx_path = os.path.splitext(video_path)[0] + ".docx"
    try:
        print("Extracting audio...")
        extract_audio(video_path, audio_file)

        print("Transcribing...")
        blocks = transcribe_with_whisper(audio_file, model_name=model_name, block_seconds=block_seconds)

        print("Saving to DOCX...")
        save_to_docx(blocks, docx_path)

        print(f"Done! Saved result to {docx_path}")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def find_videos_in_folder(folder: str) -> list[str]:
    video_files = []
    for root, _, files in os.walk(folder):
        for f in files:
            if f.lower().endswith(VIDEO_EXTENSIONS):
                video_files.append(os.path.join(root, f))
    return video_files


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Local video transcription with Whisper")
    parser.add_argument("input", help="Path to video file or folder containing videos")
    parser.add_argument("--model", default="base", help="Whisper model to use (default: base)")
    parser.add_argument("--block", type=int, default=60, help="Block length in seconds for text grouping (default: 60)")
    parser.add_argument("--device", choices=["cpu", "cuda"], default=None, help="Device to use: cpu or cuda (default: auto)")

    args = parser.parse_args()

    DEVICE = get_device(args.device)
    FFMPEG_PATH = get_ffmpeg_path()
    FFMPEG_DIR = os.path.dirname(FFMPEG_PATH)
    os.environ["PATH"] = FFMPEG_DIR + os.pathsep + os.environ["PATH"]

    input_path = args.input

    if os.path.isfile(input_path) and input_path.lower().endswith(VIDEO_EXTENSIONS):
        process_video(input_path, model_name=args.model, block_seconds=args.block)
    elif os.path.isdir(input_path):
        videos = find_videos_in_folder(input_path)
        if not videos:
            print("No video files found in the folder.")
        else:
            print(f"Found {len(videos)} video file(s).")
            for v in videos:
                process_video(v, model_name=args.model, block_seconds=args.block)
    else:
        print("Error: provide a path to a video file or folder containing videos.")
